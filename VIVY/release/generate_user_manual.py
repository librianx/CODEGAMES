from __future__ import annotations

from pathlib import Path
from html import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "VIVY_User_Manual.docx"
HTML_PATH = OUT_DIR / "VIVY_User_Manual.html"


SECTIONS = [
    (
        "一、软件简介",
        [
            "VIVY 是一款桌面陪伴型 AI 桌宠，集成了对话、记忆、创作辅助、沉浸写作、图片理解、语音输入、语音回放、本地唱歌、待机收起等功能。",
            "程序以本地桌宠界面运行，不需要单独打开网页。聊天由本地后端驱动，语音可接入 GPT-SoVITS，本地保存用户记忆数据。",
        ],
    ),
    (
        "二、启动方式",
        [
            "如果使用打包后的发布版，请先完整解压压缩包，再双击 start_all.bat。",
            "程序会先启动本地语音服务，再自动启动桌宠。",
            "首次启动时若未配置 API Key，程序会提示输入；不输入也能启动，但聊天能力会进入离线兜底模式。",
            "如果电脑性能较低，首次语音服务启动可能较慢，请耐心等待。",
        ],
    ),
    (
        "三、主界面组成",
        [
            "文字气泡区：显示 VIVY 的回答、状态提示、灵感内容等。",
            "快捷功能区：今日灵感、脑洞短剧、换个问题、形态切换。",
            "输入区：文本输入、图片上传、发送、语音输入、语音回放开关。",
            "兴趣反馈区：感兴趣、不感兴趣、清除。",
            "角色显示区：显示桌宠形象；创作模式下会出现创作领域按钮。",
            "记忆模块：可查看、编辑、保存用户记忆信息。",
        ],
    ),
    (
        "四、基础操作",
        [
            "拖动桌宠：按住桌宠窗口左键拖动，可以移动桌宠位置。",
            "双击桌宠：可在展开状态和待机收起状态之间切换。",
            "右键桌宠：弹出功能菜单，可进行更多设置和管理。",
        ],
    ),
    (
        "五、对话功能",
        [
            "文本聊天：在输入框中输入内容后点击“发送”，或直接按回车发送。",
            "普通形态下，VIVY 的回复更简短、更偏陪伴感；创作形态下，回复更发散、更适合创作交流。",
            "复制回复：点击“复制回复”可复制最近一次完整回复；若先选中了部分文字，则优先复制选中内容。",
            "换个问题：点击后，VIVY 会重新抛出一个偏好问题或引导式问题，用于更快了解用户兴趣。",
        ],
    ),
    (
        "六、首次了解你功能",
        [
            "在首次使用或某些阶段，VIVY 会主动提出偏好问题，例如喜欢的内容风格、最近关注的话题、偏好的交流方式等。",
            "当弹出问题时，选择一个选项即可。VIVY 会将回答写入记忆系统，用于优化后续回复风格与交流节奏。",
        ],
    ),
    (
        "七、快捷功能区说明",
        [
            "今日灵感：点击后生成当天的灵感分享内容，通常采用“发现 + 联想 + 邀请”的形式，适合获取创作切入点。",
            "脑洞短剧：点击后即兴生成一小段短剧，内容偏轻松、带一点反转或趣味感。",
            "形态切换：点击“形态：普通”或“形态：创作”，可在普通形态与创作形态之间切换。创作形态下会出现创作领域按钮。",
        ],
    ),
    (
        "八、创作形态功能",
        [
            "读取文档：支持 txt、md、docx、pdf。选择文档后，可输入创作目标，也可留空。VIVY 会阅读文档并以流式方式给出创作建议。",
            "创作灵感：生成一个简短、偏画面感和氛围感的创作灵感。",
            "清除参考：清除之前读取过的文档参考状态，避免当前回复继续受到旧文档影响。",
            "沉浸写作：打开独立的大屏写作窗口，适合长时间专注写作。",
        ],
    ),
    (
        "九、沉浸写作窗口说明",
        [
            "打开：载入已有的 md 或 txt 文件。",
            "最近：快速打开最近使用过的文件。",
            "新建：清空当前内容，开始新的写作任务；若有未保存内容会提示确认。",
            "保存：保存当前文件。快捷键为 Ctrl + S。",
            "另存为：将当前文本保存为新文件。",
            "恢复备份：读取自动备份文件 .vivy_immersive_autosave.md。",
            "全屏：进入或退出全屏。快捷键为 F11，Esc 可退出全屏。",
            "专注：隐藏下方辅助输出区域，让编辑界面更干净。",
            "查找：查找当前文本中的指定内容。快捷键为 Ctrl + F。",
            "目标字数：可设置目标字数，界面会显示当前进度。",
            "润色：优化表达，提高流畅度与文字美感。",
            "续写：在已有内容基础上继续发展。",
            "点评：输出对当前文本的评价、分析或结构判断。",
            "加强：强化文本的情绪、表达力度、画面感或冲突感。",
            "自定义：输入你希望 VIVY 对当前选区或全文执行的操作。",
            "复制输出：将辅助输出区结果复制到剪贴板。",
            "插入到光标：将辅助结果插入当前光标位置。",
            "替换选区：用辅助结果直接替换选中的原文；若未选中，则按插入处理。",
        ],
    ),
    (
        "十、图片理解功能",
        [
            "点击图片按钮后选择一张图片，再输入问题并发送，VIVY 会将图片和文字一起发送给支持视觉的模型。",
            "适用于图片理解、识别、描述、分析等场景。",
            "如果没有配置支持视觉的模型，图片理解可能无法正常工作。",
            "图片大小受系统设置限制，超出限制时程序会提示失败。",
            "点击图片清除按钮可取消当前已选择的图片。",
        ],
    ),
    (
        "十一、语音功能",
        [
            "语音输入：点击“语音”按钮后对着麦克风说话，识别完成后会自动填入输入框并直接发送。",
            "语音输入依赖 Windows 系统语音识别，需要麦克风和系统语音环境正常可用。",
            "语音回放开关：点击“语音回放：开/关”可控制 VIVY 的回答是否自动朗读。",
        ],
    ),
    (
        "十二、兴趣反馈功能",
        [
            "感兴趣：告诉 VIVY 你希望继续深入当前话题。",
            "不感兴趣：告诉 VIVY 当前方向不适合你，VIVY 会更倾向于切换方向。",
            "清除：清除当前兴趣信号，恢复中性状态。",
        ],
    ),
    (
        "十三、记忆模块说明",
        [
            "记忆模块默认可隐藏，右键菜单可显示或隐藏。",
            "刷新记忆：从本地数据库拉取最新记忆内容，建议编辑前先刷新。",
            "保存记忆：将修改后的 summary、summary_long、preferences JSON 写回本地数据库。",
            "操作指南：弹出一个简易说明，指导如何正确编辑记忆模块。",
            "填入示例 JSON：自动生成一个可参考的 preferences JSON 模板。",
            "summary：用于一句话概括最近状态。",
            "summary_long：用于记录长期偏好、目标、边界、稳定倾向。",
            "preferences JSON：保存结构化偏好，必须是合法 JSON。",
            "最近对话回合：显示最近一段时间的对话记录，并附带回合 ID。",
            "删除回合：输入回合 ID 后点击“删除回合”，可删除指定的一条历史对话记录。",
        ],
    ),
    (
        "十四、右键菜单功能",
        [
            "重置本机用户 ID：重新生成新的本机用户标识，相当于重新开始一套新的用户记忆。",
            "设置 API Key：输入并保存 DeepSeek API Key，会自动写入 .env。",
            "显示或隐藏记忆模块：控制右侧记忆模块的显示状态。",
            "切换待机收起：手动切换是否进入待机收起状态。",
            "设置待机时长：设置多少秒无操作后自动进入待机状态。",
            "读取文档（创作辅助）：与创作形态中的读取文档功能相同。",
            "清除已读取文档：清除当前文档参考状态。",
            "沉浸写作窗口：打开沉浸写作独立窗口。",
            "停止唱歌：停止当前本地歌曲播放。",
            "退出 VIVY：关闭程序。",
        ],
    ),
    (
        "十五、待机收起功能",
        [
            "当用户在设定时间内无操作时，VIVY 会自动进入待机收起状态。",
            "用户也可以通过右键菜单的“切换待机收起”或双击桌宠来手动切换状态。",
            "待机时长可以自行设置，常用交互建议设置长一些，希望桌面更整洁则可设置短一些。",
        ],
    ),
    (
        "十六、本地唱歌功能",
        [
            "VIVY 支持根据关键词触发本地唱歌，会从 song 文件夹中随机选择一首 wav 歌曲播放。",
            "常见触发语句包括：唱歌、来一首歌、给我唱、你唱、哄我、安慰我、我心情不好、我很难过。",
            "停止方式包括输入“停止唱歌”“停歌”“停止播放”等指令，或在右键菜单中点击“停止唱歌”。",
            "当前仅支持 wav 格式歌曲。若 song 目录为空或文件格式不符合要求，则无法播放。",
        ],
    ),
    (
        "十七、图片与语音结合使用建议",
        [
            "图片上传后，最好再配合文字提问，效果通常更好。",
            "语音输入适合快速日常交流。",
            "语音回放适合陪伴感场景；在公共环境或安静环境可关闭。",
            "创作模式下更适合处理长内容、文档、灵感与写作需求。",
        ],
    ),
    (
        "十八、数据保存位置说明",
        [
            "vivy.sqlite：本地数据库，保存记忆与部分交互记录。",
            ".desktop_user_id：本机用户 ID。",
            ".vivy_immersive_autosave.md：沉浸写作自动备份。",
            ".immersive_recent.json：最近打开文件记录。",
            ".env：程序配置项，例如 API Key、待机时长、语音参数等。",
        ],
    ),
    (
        "十九、常见问题",
        [
            "启动后没有对话能力：通常是没有配置 API Key，可通过右键菜单的“设置 API Key”补充。",
            "图片功能不能用：通常是当前模型不支持视觉，需配置支持图像的模型。",
            "语音输入失败：通常与麦克风权限或 Windows 语音识别环境有关。",
            "不能唱歌：通常是 song 文件夹里没有可用的 wav 文件，或文件损坏。",
            "沉浸写作辅助没有反应：通常是当前没有可处理内容，或网络 / 模型请求失败。",
        ],
    ),
    (
        "二十、推荐使用流程",
        [
            "日常陪伴模式：启动程序后直接聊天，按需要开启语音回放，并使用“今日灵感”或“脑洞短剧”放松交流。",
            "创作辅助模式：切换到创作形态，读取文档，使用创作灵感，再进入沉浸写作窗口深入处理。",
            "深度写作模式：打开沉浸写作，设置目标字数，边写边使用润色、续写、加强、点评等功能，并及时保存。",
        ],
    ),
]


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def configure_doc_style(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.2)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("VIVY 桌宠使用说明书")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x19, 0x3A, 0x63)
    set_east_asia_font(run, "Microsoft YaHei")

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(8)
    run = p2.add_run("适用于桌宠正式发布版")
    run.italic = True
    run.font.size = Pt(12)
    set_east_asia_font(run, "Microsoft YaHei")

    p3 = document.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(120)
    run = p3.add_run("编制日期：2026 年 4 月")
    run.font.size = Pt(11)
    set_east_asia_font(run, "Microsoft YaHei")

    document.add_section(WD_SECTION.NEW_PAGE)


def build_docx() -> None:
    document = Document()
    configure_doc_style(document)
    add_cover(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("正文")
    run.bold = True
    run.font.size = Pt(16)
    set_east_asia_font(run, "Microsoft YaHei")

    intro = document.add_paragraph()
    intro.paragraph_format.space_after = Pt(10)
    run = intro.add_run(
        "本说明书用于帮助用户快速了解 VIVY 桌宠的各项功能、常见操作方法及使用建议。"
    )
    set_east_asia_font(run, "Microsoft YaHei")

    for heading, items in SECTIONS:
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(heading)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1D, 0x4E, 0x89)
        set_east_asia_font(run, "Microsoft YaHei")

        for item in items:
            bp = document.add_paragraph(style="List Bullet")
            bp.paragraph_format.space_after = Pt(2)
            run = bp.add_run(item)
            set_east_asia_font(run, "Microsoft YaHei")

    document.save(DOCX_PATH)


def build_html() -> None:
    style = """
    body {
      font-family: "Microsoft YaHei", "PingFang SC", sans-serif;
      color: #17212b;
      line-height: 1.7;
      margin: 0;
      background: #f4f7fb;
    }
    .page {
      width: 860px;
      margin: 24px auto;
      background: #ffffff;
      box-shadow: 0 10px 30px rgba(0,0,0,0.08);
      padding: 54px 64px;
    }
    h1 {
      text-align: center;
      color: #183b63;
      font-size: 28px;
      margin: 18px 0 8px;
    }
    .subtitle {
      text-align: center;
      color: #5f7388;
      font-size: 14px;
      margin-bottom: 36px;
    }
    h2 {
      color: #1d4e89;
      font-size: 20px;
      margin: 26px 0 10px;
      border-bottom: 1px solid #d7e2f0;
      padding-bottom: 6px;
    }
    ul {
      margin: 0 0 8px 0;
      padding-left: 24px;
    }
    li {
      margin: 7px 0;
    }
    .intro {
      margin-bottom: 16px;
    }
    """
    parts = [
        "<!doctype html>",
        "<html><head><meta charset='utf-8'>",
        f"<style>{style}</style>",
        "</head><body>",
        "<div class='page'>",
        "<h1>VIVY 桌宠使用说明书</h1>",
        "<div class='subtitle'>适用于桌宠正式发布版</div>",
        "<p class='intro'>本说明书用于帮助用户快速了解 VIVY 桌宠的各项功能、常见操作方法及使用建议。</p>",
    ]
    for heading, items in SECTIONS:
        parts.append(f"<h2>{escape(heading)}</h2>")
        parts.append("<ul>")
        for item in items:
            parts.append(f"<li>{escape(item)}</li>")
        parts.append("</ul>")
    parts.extend(["</div>", "</body></html>"])
    HTML_PATH.write_text("\n".join(parts), encoding="utf-8")


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_docx()
    build_html()
    print(DOCX_PATH)
    print(HTML_PATH)


if __name__ == "__main__":
    main()
