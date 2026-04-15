from __future__ import annotations

from html import escape
from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "release"
ASSET_DIR = OUT_DIR / "tutorial_assets"
DOCX_PATH = OUT_DIR / "VIVY_User_Tutorial_v2.docx"
HTML_PATH = OUT_DIR / "VIVY_User_Tutorial_v2.html"

BG = (8, 14, 23)
PANEL = (20, 31, 46)
PANEL_ALT = (26, 42, 61)
PANEL_3 = (15, 45, 52)
ACCENT = (55, 196, 235)
ACCENT_SOFT = (132, 234, 255)
GREEN = (93, 210, 157)
WARN = (255, 190, 100)
TEXT = (239, 248, 255)
TEXT_DIM = (185, 205, 224)
LINE = (74, 132, 156)


SECTIONS = [
    {
        "title": "1. 这份教程适合谁",
        "body": [
            "适合第一次拿到 VIVY 压缩包的用户，也适合需要给别人演示或分发 VIVY 的人。",
            "教程按“先跑起来，再会使用，再懂设置，再懂数据安全”的顺序编写。",
            "如果只是普通使用，重点看第 2 到第 8 节；如果要打包给别人，重点看第 11 节。",
        ],
    },
    {
        "title": "2. 快速启动",
        "diagram": "startup_flow.png",
        "body": [
            "Windows：完整解压压缩包后，进入 Windows\\VIVY_bundle，双击 start_all.bat。脚本会先启动本地语音服务，再启动桌宠。",
            "macOS：进入 macOS 文件夹，双击 start_all.command；如果系统拦截，右键该文件后选择“打开”。脚本会进入 VIVY-src 并准备 Python 环境。",
            "首次启动可能较慢，尤其是需要安装依赖或启动本地语音服务时，请等待窗口出现。",
        ],
        "tips": [
            "不要直接在压缩包里运行，先完整解压。",
            "如果双击没有反应，可以用终端/PowerShell 运行 start_all 脚本查看报错。",
        ],
    },
    {
        "title": "3. 第一次配置 API Key",
        "body": [
            "首次启动如果没有 .env 或没有 API Key，VIVY 会提示你输入。",
            "也可以右键桌宠，点击“设置 API Key”，输入 DeepSeek 或 OpenAI 兼容服务的 API Key。",
            "API Key 会保存到 .env。这个文件是私人配置，打包给别人时不要一起发。",
        ],
        "examples": [
            "如果只是测试界面，不填 API Key 也能启动，但聊天能力可能进入离线兜底。",
            "如果要使用图片理解，需要额外配置支持视觉的模型。普通文字模型不一定能看图。",
        ],
    },
    {
        "title": "4. 主界面怎么用",
        "diagram": "interface_map.png",
        "body": [
            "左侧气泡区显示 Vivy 的回复、状态提示、天气问候和灵感内容。",
            "底部输入区可以输入文字，按回车或点击“发送”提交；图片按钮可附加图片；语音按钮可启动语音输入。",
            "快捷按钮包括今日灵感、脑洞短剧、换个问题、形态切换、兴趣反馈等。",
            "双击桌宠可以展开或收起；按住左键拖动可以移动位置；右键打开设置菜单。",
        ],
    },
    {
        "title": "5. 日常聊天与本地快捷能力",
        "diagram": "message_flow.png",
        "body": [
            "VIVY 会先判断本地能力：时间查询、天气查询、唱歌/停止唱歌等会直接处理，不必请求大模型。",
            "普通聊天、创作交流、复杂问题会进入后端 API，读取记忆和运行时上下文后再请求模型。",
            "新的问题会中断旧回复和旧语音，避免多个回答串在一起。",
        ],
        "examples": [
            "问时间：现在几点？今天几号？",
            "问天气：今天会下雨吗？上海明天几度？",
            "普通聊天：帮我想一个故事开头。",
        ],
    },
    {
        "title": "6. 天气、位置与每日问候",
        "diagram": "daily_greeting_flow.png",
        "body": [
            "右键桌宠可以设置当前位置。这个城市会作为天气查询和早晨问候的默认城市。",
            "右键菜单可以开启或关闭 IP 粗定位。IP 粗定位只用于天气，不做后台持续定位。",
            "每天第一次打开会触发一次问候；早、中、晚、深夜也有自动问候时间点。",
            "VIVY 会把当天打开次数和已问候时段写入记忆，因此关闭后重新打开不会重复第一次问候。",
        ],
        "tips": [
            "如果不想自动问候，可以右键关闭“每日问候”。",
            "早晨问候会包含天气和出门防护提醒；其他时段主要是陪伴式问候。",
        ],
    },
    {
        "title": "7. 情绪陪伴与唱歌",
        "body": [
            "现在 VIVY 不会因为你说“难过、心情不好、想哭”就立刻唱歌。",
            "轻度低落时，VIVY 会先安慰；明显难过时，VIVY 会先陪伴并询问要不要听歌。",
            "只有你明确说“唱首歌、想听你唱、唱吧、好”等，才会播放 song 文件夹里的 WAV 歌曲。",
            "如果说“停止唱歌、别唱了、停歌”，VIVY 会停止当前歌曲。",
        ],
        "examples": [
            "不会直接唱：我今天好难过。",
            "会进入邀请：我有点想哭。",
            "会唱：唱首歌给我听。 / 唱吧。",
        ],
    },
    {
        "title": "8. 创作模式与沉浸写作",
        "body": [
            "点击“形态：普通/创作”可以在普通陪伴模式和创作模式之间切换。",
            "创作模式下可以读取 txt、md、docx、pdf 文件，让 Vivy 基于文档进行分析或辅助创作。",
            "右键菜单或创作按钮可以打开“沉浸写作窗口”，适合长文本写作。",
            "沉浸写作支持新建、打开、保存、全屏、目标字数、润色、续写、点评、加强和自定义指令。",
        ],
        "tips": [
            "读取文档前，建议先确认文档不是过大的扫描版 PDF。",
            "长文写作时记得定期保存；自动备份文件是 .vivy_immersive_autosave.md。",
        ],
    },
    {
        "title": "9. 记忆模块怎么理解",
        "diagram": "memory_model.png",
        "body": [
            "VIVY 的记忆主要保存在 vivy.sqlite 中，包含用户偏好、摘要、最近对话回合和每日问候状态。",
            "记忆面板里的 preferences JSON 是结构化偏好，必须保持合法 JSON，否则保存会失败。",
            "daily_greeting_state 会记录每天 open_count、open_greeted、slots 等状态，用来避免重复问候。",
            "重置本机用户 ID 会开启一套新的本机身份，相当于从新的记忆开始。",
        ],
    },
    {
        "title": "10. 右键菜单速查",
        "body": [
            "设置 API Key：修改模型接口凭证。",
            "设置当前位置：修改天气默认城市。",
            "开启/关闭 IP 粗定位：控制是否允许用 IP 推断城市。",
            "开启/关闭每日问候：控制自动问候功能。",
            "显示/隐藏记忆模块：控制记忆编辑面板。",
            "设置待机时长：设置无操作多久后收起。",
            "读取文档、清除参考、沉浸写作窗口：创作辅助入口。",
            "停止唱歌、退出 VIVY：控制播放和关闭程序。",
        ],
    },
    {
        "title": "11. 打包给别人时怎么保持干净",
        "diagram": "distribution_safety.png",
        "body": [
            "如果把 VIVY 压缩包发给别人，建议不要带自己的记忆文件和私人配置。",
            "必须排除：vivy.sqlite、.desktop_user_id、.env。",
            "建议排除：.vivy_immersive_autosave.md、.immersive_recent.json。",
            "可以保留：env.example、song 文件夹、static 资源、启动脚本、exe 或 macOS 源码启动包。",
            "这样别人第一次运行会生成新的用户 ID 和新的数据库，不会继承你的记忆。",
        ],
    },
    {
        "title": "12. 常见问题排查",
        "diagram": "troubleshooting_tree.png",
        "body": [
            "启动失败：先确认是否完整解压，再用 start_all 脚本在终端中运行查看错误。",
            "不能聊天：检查 API Key 是否配置，网络是否可访问模型服务。",
            "天气查不到：先右键设置当前位置，城市名尽量写清楚，例如“上海市”“茂名市”。",
            "没有自动问候：检查每日问候开关；当天已问候过也不会重复播报。",
            "不能唱歌：确认 song 文件夹里有可播放的 WAV 文件。",
            "图片理解无效：确认配置了支持视觉的模型。",
        ],
    },
]


def font_path() -> str | None:
    candidates = [
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
        Path("/System/Library/Fonts/PingFang.ttc"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return str(path)
    return None


FONT_PATH = font_path()


def pil_font(size: int, bold: bool = False):
    if FONT_PATH:
        return ImageFont.truetype(FONT_PATH, size=size)
    return ImageFont.load_default()


def text_size(draw: ImageDraw.ImageDraw, text: str, font) -> tuple[int, int]:
    if not text:
        return (0, 0)
    box = draw.textbbox((0, 0), text, font=font)
    return (box[2] - box[0], box[3] - box[1])


def wrap_px(draw: ImageDraw.ImageDraw, text: str, font, max_w: int) -> list[str]:
    lines: list[str] = []
    for raw in text.split("\n"):
        line = ""
        for ch in raw:
            test = line + ch
            if text_size(draw, test, font)[0] <= max_w or not line:
                line = test
            else:
                lines.append(line)
                line = ch
        if line:
            lines.append(line)
    return lines or [""]


def draw_multiline(draw, xy, text, font, fill, max_w, line_gap=6):
    x, y = xy
    for line in wrap_px(draw, text, font, max_w):
        draw.text((x, y), line, font=font, fill=fill)
        y += text_size(draw, line, font)[1] + line_gap
    return y


def box(draw, xywh, title, body="", fill=PANEL, outline=ACCENT, title_color=ACCENT_SOFT):
    x, y, w, h = xywh
    draw.rounded_rectangle((x, y, x + w, y + h), radius=22, fill=fill, outline=outline, width=3)
    title_font = pil_font(28)
    body_font = pil_font(22)
    draw_multiline(draw, (x + 24, y + 18), title, title_font, title_color, w - 48, line_gap=8)
    if body:
        draw_multiline(draw, (x + 24, y + 62), body, body_font, TEXT, w - 48, line_gap=8)


def arrow(draw, start, end, fill=ACCENT):
    draw.line((start, end), fill=fill, width=5)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) > abs(ey - sy):
        direction = 1 if ex > sx else -1
        pts = [(ex, ey), (ex - direction * 18, ey - 10), (ex - direction * 18, ey + 10)]
    else:
        direction = 1 if ey > sy else -1
        pts = [(ex, ey), (ex - 10, ey - direction * 18), (ex + 10, ey - direction * 18)]
    draw.polygon(pts, fill=fill)


def canvas(title: str):
    img = Image.new("RGB", (1500, 850), BG)
    draw = ImageDraw.Draw(img)
    draw.text((60, 42), title, font=pil_font(38), fill=TEXT)
    draw.rounded_rectangle((60, 102, 1440, 108), radius=3, fill=ACCENT)
    return img, draw


def save_img(img: Image.Image, name: str):
    path = ASSET_DIR / name
    img.save(path)
    return path


def diagram_startup():
    img, draw = canvas("图 1：启动流程")
    items = [
        ("完整解压", "不要直接在压缩包中运行"),
        ("启动脚本", "Windows: start_all.bat\nmacOS: start_all.command"),
        ("准备环境", "启动语音服务或创建 venv"),
        ("启动桌宠", "VIVY.exe / desktop_pet_dual_platform.py"),
        ("进入使用", "配置 API Key 后开始聊天"),
    ]
    x = 95
    y = 205
    w = 235
    for i, (t, b) in enumerate(items):
        box(draw, (x + i * 275, y, w, 205), t, b, fill=PANEL_ALT if i % 2 else PANEL)
        if i < len(items) - 1:
            arrow(draw, (x + i * 275 + w + 12, y + 102), (x + (i + 1) * 275 - 12, y + 102))
    box(draw, (230, 520, 1040, 150), "首次启动提示", "如果没有 .env 或 API Key，VIVY 会提示输入。普通用户推荐通过右键菜单设置 API Key。", fill=PANEL_3, outline=GREEN)
    return save_img(img, "startup_flow.png")


def diagram_interface():
    img, draw = canvas("图 2：主界面功能分区")
    box(draw, (90, 170, 570, 220), "回复气泡区", "显示回答、天气问候、状态提示、灵感内容。")
    box(draw, (90, 430, 570, 120), "快捷按钮区", "今日灵感 / 脑洞短剧 / 换个问题 / 形态切换 / 兴趣反馈", fill=PANEL_ALT)
    box(draw, (90, 590, 570, 120), "输入与语音区", "输入文字、选择图片、发送、语音输入、语音回放开关。")
    box(draw, (760, 185, 520, 370), "角色显示区", "显示 VIVY 形象；可拖动移动，双击收起或展开。", fill=PANEL_3)
    box(draw, (760, 600, 520, 110), "右键菜单", "API Key、当前位置、每日问候、记忆模块、待机时长、退出。", fill=PANEL_ALT, outline=WARN)
    arrow(draw, (660, 275), (760, 340), fill=ACCENT_SOFT)
    arrow(draw, (660, 655), (760, 655), fill=WARN)
    return save_img(img, "interface_map.png")


def diagram_message():
    img, draw = canvas("图 3：一次消息处理流程")
    rows = [
        ("用户输入", "文字 / 语音 / 图片 / 快捷按钮"),
        ("本地分流", "时间、天气、唱歌、停止唱歌优先处理"),
        ("后端 API", "读取记忆、偏好、运行时上下文"),
        ("模型回复", "普通回复或流式回复"),
        ("前端呈现", "气泡显示、语音播报、刷新记忆"),
    ]
    y = 170
    for i, (t, b) in enumerate(rows):
        box(draw, (315, y + i * 120, 870, 78), t, b, fill=PANEL_ALT if i % 2 else PANEL)
        if i < len(rows) - 1:
            arrow(draw, (750, y + i * 120 + 82), (750, y + (i + 1) * 120 - 6))
    box(draw, (80, 360, 210, 160), "中断机制", "新问题会中断旧回复和旧语音。", fill=PANEL_3, outline=GREEN)
    box(draw, (1210, 360, 210, 160), "本地优先", "简单能力无需请求大模型。", fill=PANEL_3, outline=GREEN)
    return save_img(img, "message_flow.png")


def diagram_greeting():
    img, draw = canvas("图 4：每日问候与打开次数记忆")
    box(draw, (100, 165, 340, 145), "每次启动", "生成 open_event_id\n调用 trigger=open")
    box(draw, (580, 165, 340, 145), "读取记忆", "daily_greeting_state\nopen_count / slots")
    box(draw, (1060, 165, 340, 145), "判断是否问候", "首次打开才说\n重开不重复")
    arrow(draw, (440, 238), (580, 238))
    arrow(draw, (920, 238), (1060, 238))

    box(draw, (100, 455, 340, 145), "定时触发", "早 / 中 / 晚 / 深夜\n调用 trigger=schedule", fill=PANEL_ALT)
    box(draw, (580, 455, 340, 145), "检查 slots", "该时段是否已经问候", fill=PANEL_ALT)
    box(draw, (1060, 455, 340, 145), "写回状态", "完成后保存 slot\n保留最近 14 天", fill=PANEL_ALT)
    arrow(draw, (440, 528), (580, 528), fill=GREEN)
    arrow(draw, (920, 528), (1060, 528), fill=GREEN)

    draw.text((195, 690), "结果：关闭 VIVY 后当天重新打开，只增加 open_count，不重复首次问候。", font=pil_font(26), fill=ACCENT_SOFT)
    return save_img(img, "daily_greeting_flow.png")


def diagram_memory():
    img, draw = canvas("图 5：本地记忆模型")
    box(draw, (120, 180, 360, 150), "vivy.sqlite", "用户偏好、摘要、对话回合、每日问候状态。", fill=PANEL_3, outline=GREEN)
    box(draw, (570, 180, 360, 150), ".desktop_user_id", "本机用户身份；重置后会生成新的记忆线。", fill=PANEL_ALT)
    box(draw, (1020, 180, 360, 150), ".env", "API Key、天气和语音等配置。", fill=PANEL_ALT, outline=WARN)
    box(draw, (350, 500, 800, 145), "记忆面板", "可以查看 summary、summary_long、preferences JSON 和最近对话回合。\n保存前请确保 JSON 合法。")
    arrow(draw, (300, 330), (520, 500), fill=GREEN)
    arrow(draw, (750, 330), (750, 500))
    arrow(draw, (1200, 330), (980, 500), fill=WARN)
    return save_img(img, "memory_model.png")


def diagram_distribution():
    img, draw = canvas("图 6：分发给别人时的文件安全")
    box(draw, (110, 180, 500, 370), "可以放进压缩包", "VIVY.exe / start_all 脚本\nstatic 资源\nsong 文件夹\nenv.example\nmacOS 的 VIVY-src", fill=PANEL_3, outline=GREEN)
    box(draw, (890, 180, 500, 370), "不要放进压缩包", "vivy.sqlite\n.desktop_user_id\n.env\n.vivy_immersive_autosave.md\n.immersive_recent.json", fill=PANEL_ALT, outline=WARN)
    arrow(draw, (610, 360), (890, 360), fill=ACCENT)
    draw.text((270, 650), "这样别人首次运行会生成自己的用户 ID 和数据库，不会继承你的记忆或 API Key。", font=pil_font(26), fill=ACCENT_SOFT)
    return save_img(img, "distribution_safety.png")


def diagram_troubleshooting():
    img, draw = canvas("图 7：常见问题排查路径")
    box(draw, (555, 150, 390, 80), "遇到问题", "先确认完整解压并从 start_all 启动")
    box(draw, (110, 330, 300, 135), "不能聊天", "检查 API Key\n检查网络和模型服务", fill=PANEL_ALT)
    box(draw, (445, 330, 300, 135), "天气查不到", "右键设置当前位置\n城市名写清楚", fill=PANEL_ALT)
    box(draw, (780, 330, 300, 135), "不能唱歌", "song 文件夹需要 WAV\n确认没有停止播放", fill=PANEL_ALT)
    box(draw, (1115, 330, 300, 135), "没有问候", "确认每日问候开启\n当天已问候不会重复", fill=PANEL_ALT)
    for x in (260, 595, 930, 1265):
        arrow(draw, (750, 230), (x, 330), fill=WARN)
    box(draw, (380, 600, 740, 115), "还不行怎么办", "用 PowerShell / 终端运行 start_all，截图错误信息，优先看最后几行。", fill=PANEL_3, outline=GREEN)
    return save_img(img, "troubleshooting_tree.png")


def build_diagrams():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    return {
        "startup_flow.png": diagram_startup(),
        "interface_map.png": diagram_interface(),
        "message_flow.png": diagram_message(),
        "daily_greeting_flow.png": diagram_greeting(),
        "memory_model.png": diagram_memory(),
        "distribution_safety.png": diagram_distribution(),
        "troubleshooting_tree.png": diagram_troubleshooting(),
    }


def set_east_asia_font(run, font_name: str = "Microsoft YaHei") -> None:
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def style_doc(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16 if level == 1 else 13)
    run.font.color.rgb = RGBColor(0x1D, 0x4E, 0x89)
    set_east_asia_font(run)
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        set_east_asia_font(run)


def add_examples(doc: Document, title: str, items: list[str]):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(0x1D, 0x4E, 0x89)
    set_east_asia_font(r)
    add_bullets(doc, items)


def build_docx(diagrams: dict[str, Path]):
    doc = Document()
    style_doc(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(58)
    run = p.add_run("VIVY 桌宠使用教程")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x18, 0x3B, 0x63)
    set_east_asia_font(run)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run("快速启动｜界面操作｜天气问候｜记忆说明｜打包分发")
    run.font.size = Pt(12)
    run.italic = True
    run.font.color.rgb = RGBColor(0x5F, 0x73, 0x88)
    set_east_asia_font(run)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(150)
    r3 = p3.add_run("版本：VIVY 双平台版    日期：2026 年 4 月")
    r3.font.size = Pt(11)
    set_east_asia_font(r3)
    doc.add_page_break()

    add_heading(doc, "教程目录")
    for sec in SECTIONS:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(sec["title"])
        set_east_asia_font(r)
    doc.add_page_break()

    for sec in SECTIONS:
        add_heading(doc, sec["title"])
        if sec.get("diagram"):
            doc.add_picture(str(diagrams[sec["diagram"]]), width=Cm(16.3))
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cap.add_run(sec["diagram"].replace("_", " ").replace(".png", ""))
            cr.font.size = Pt(9)
            cr.font.color.rgb = RGBColor(0x67, 0x7D, 0x92)
            set_east_asia_font(cr)
        add_bullets(doc, sec["body"])
        if sec.get("tips"):
            add_examples(doc, "使用提示", sec["tips"])
        if sec.get("examples"):
            add_examples(doc, "常见说法 / 示例", sec["examples"])

    add_heading(doc, "附录：建议保留与排除的文件")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["类别", "文件", "说明"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    rows = [
        ("可保留", "env.example", "示例配置，不包含私人 API Key。"),
        ("可保留", "song 文件夹", "放 WAV 歌曲，用于唱歌功能。"),
        ("可保留", "static 资源", "桌宠头像和界面资源。"),
        ("应排除", ".env", "包含 API Key 等私人配置。"),
        ("应排除", "vivy.sqlite", "包含你的记忆、偏好和对话记录。"),
        ("应排除", ".desktop_user_id", "你的本机用户 ID。"),
        ("建议排除", ".vivy_immersive_autosave.md", "沉浸写作自动备份。"),
        ("建议排除", ".immersive_recent.json", "最近打开文件记录。"),
    ]
    for kind, file, desc in rows:
        row = table.add_row().cells
        row[0].text = kind
        row[1].text = file
        row[2].text = desc

    doc.save(DOCX_PATH)


def build_html(diagrams: dict[str, Path]):
    css = dedent(
        """
        body { margin: 0; background: #f3f7fb; color: #16222d; font-family: "Microsoft YaHei", "PingFang SC", sans-serif; line-height: 1.72; }
        .page { width: 920px; margin: 26px auto; padding: 52px 66px; background: white; box-shadow: 0 10px 28px rgba(12, 28, 45, .12); }
        h1 { text-align: center; color: #183b63; margin: 0; font-size: 32px; }
        .sub { text-align: center; color: #66798c; margin: 10px 0 36px; }
        h2 { color: #1d4e89; margin-top: 32px; padding-bottom: 6px; border-bottom: 1px solid #d8e3ef; }
        ul { padding-left: 23px; }
        li { margin: 6px 0; }
        img { display: block; max-width: 100%; margin: 14px auto 8px; border-radius: 8px; box-shadow: 0 8px 20px rgba(18, 36, 54, .18); }
        .caption { text-align: center; color: #6a7e91; font-size: 13px; margin-bottom: 10px; }
        .box { background: #f4f8fc; border-left: 4px solid #1d9ed0; padding: 10px 14px; margin: 12px 0; }
        table { width: 100%; border-collapse: collapse; margin: 16px 0; }
        th, td { border: 1px solid #cbd8e5; padding: 8px 10px; text-align: left; }
        th { background: #eaf3fb; color: #183b63; }
        """
    )
    parts = [
        "<!doctype html><html><head><meta charset='utf-8'>",
        f"<style>{css}</style></head><body><div class='page'>",
        "<h1>VIVY 桌宠使用教程</h1>",
        "<div class='sub'>快速启动｜界面操作｜天气问候｜记忆说明｜打包分发</div>",
    ]
    for sec in SECTIONS:
        parts.append(f"<h2>{escape(sec['title'])}</h2>")
        if sec.get("diagram"):
            src = f"tutorial_assets/{sec['diagram']}"
            parts.append(f"<img src='{escape(src)}' alt='{escape(sec['title'])}'>")
            parts.append(f"<div class='caption'>{escape(sec['diagram'])}</div>")
        parts.append("<ul>")
        for item in sec["body"]:
            parts.append(f"<li>{escape(item)}</li>")
        parts.append("</ul>")
        if sec.get("tips"):
            parts.append("<div class='box'><strong>使用提示</strong><ul>")
            for item in sec["tips"]:
                parts.append(f"<li>{escape(item)}</li>")
            parts.append("</ul></div>")
        if sec.get("examples"):
            parts.append("<div class='box'><strong>常见说法 / 示例</strong><ul>")
            for item in sec["examples"]:
                parts.append(f"<li>{escape(item)}</li>")
            parts.append("</ul></div>")

    parts.append("<h2>附录：建议保留与排除的文件</h2>")
    parts.append("<table><tr><th>类别</th><th>文件</th><th>说明</th></tr>")
    for row in [
        ("可保留", "env.example", "示例配置，不包含私人 API Key。"),
        ("可保留", "song 文件夹", "放 WAV 歌曲，用于唱歌功能。"),
        ("应排除", ".env", "包含 API Key 等私人配置。"),
        ("应排除", "vivy.sqlite", "包含你的记忆、偏好和对话记录。"),
        ("应排除", ".desktop_user_id", "你的本机用户 ID。"),
        ("建议排除", ".vivy_immersive_autosave.md", "沉浸写作自动备份。"),
        ("建议排除", ".immersive_recent.json", "最近打开文件记录。"),
    ]:
        parts.append("<tr>" + "".join(f"<td>{escape(x)}</td>" for x in row) + "</tr>")
    parts.append("</table></div></body></html>")
    HTML_PATH.write_text("\n".join(parts), encoding="utf-8")


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = build_diagrams()
    build_docx(diagrams)
    build_html(diagrams)
    print(DOCX_PATH)
    print(HTML_PATH)
    for path in diagrams.values():
        print(path)


if __name__ == "__main__":
    main()
