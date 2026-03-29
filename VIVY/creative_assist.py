from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any, Optional


@dataclass
class LoadedDocument:
    path: str
    ext: str
    text: str


def load_document_text(path: str, max_chars: int = 60_000) -> LoadedDocument:
    p = Path(path)
    ext = p.suffix.lower().lstrip(".")
    if ext in ("txt", "md"):
        text = p.read_text(encoding="utf-8", errors="ignore")
        return LoadedDocument(path=str(p), ext=ext, text=text[:max_chars])

    if ext == "docx":
        try:
            from docx import Document  # type: ignore
        except Exception as e:
            raise RuntimeError("缺少依赖 python-docx，无法读取 .docx") from e
        doc = Document(str(p))
        parts = []
        for para in doc.paragraphs:
            t = (para.text or "").strip()
            if t:
                parts.append(t)
        text = "\n".join(parts)
        return LoadedDocument(path=str(p), ext=ext, text=text[:max_chars])

    if ext == "pdf":
        try:
            from pypdf import PdfReader  # type: ignore
        except Exception as e:
            raise RuntimeError("缺少依赖 pypdf，无法读取 .pdf") from e
        reader = PdfReader(str(p))
        parts = []
        for page in reader.pages:
            try:
                t = page.extract_text() or ""
            except Exception:
                t = ""
            t = t.strip()
            if t:
                parts.append(t)
        text = "\n\n".join(parts)
        return LoadedDocument(path=str(p), ext=ext, text=text[:max_chars])

    raise RuntimeError(f"不支持的文件类型：.{ext}（支持 txt/md/docx/pdf）")


# WPS/Word 加载项或 COM 桥：选区辅助（短请求，适合边写边点）
OFFICE_PASSAGE_MAX = 14_000
OFFICE_CONTEXT_MAX = 5_000
# 沉浸写作 / Office 辅助：用户投喂的参考文档总上限（与 passage 分开计费）
OFFICE_REFERENCE_MAX_PER_DOC = 10_000
OFFICE_REFERENCE_MAX_TOTAL = 22_000

def normalize_office_reference_docs(raw: Any) -> list[tuple[str, str]]:
    """解析客户端 reference_docs，按总字数预算截断。每项为 (显示名, 正文)。"""
    out: list[tuple[str, str]] = []
    if not isinstance(raw, list):
        return out
    budget = OFFICE_REFERENCE_MAX_TOTAL
    for item in raw:
        if not isinstance(item, dict):
            continue
        label = str(item.get("label") or item.get("name") or "参考文档").strip() or "参考文档"
        text = str(item.get("text") or "").strip()
        if not text:
            continue
        cap = min(OFFICE_REFERENCE_MAX_PER_DOC, max(0, budget))
        if cap <= 0:
            break
        if len(text) > cap:
            ellip = "\n…（已截断）"
            take = max(0, cap - len(ellip))
            text = text[:take] + ellip
        label = label[:200]
        out.append((label, text))
        budget -= len(text)
        if budget <= 0:
            break
    return out


OFFICE_ACTION_HINTS = {
    "polish": "请只针对【选中文本】做润色：保留原意与人称，输出一版可直接替换的正文（不要前言后语）。",
    "continue": "请根据【选中文本】自然续写 2～6 句，风格一致，直接输出续写正文。",
    "critique": "请点评【选中文本】：问题、优点各若干条，再给 2 条可执行修改建议（简练）。",
    "improve": "请在保持情节/论点不变的前提下，加强画面感或逻辑衔接；给出「修改后完整选段」替换原文。",
    "free": "",  # 仅使用 user_goal
}


def build_office_passage_prompt(
    passage: str,
    action: str = "polish",
    user_goal: Optional[str] = None,
    context_excerpt: Optional[str] = None,
    reference_docs: Optional[list[tuple[str, str]]] = None,
) -> str:
    passage = (passage or "").strip()
    if not passage:
        raise ValueError("选区文本为空")
    goal = (user_goal or "").strip()
    ctx = (context_excerpt or "").strip()
    act = (action or "polish").strip().lower()
    hint = OFFICE_ACTION_HINTS.get(act) or OFFICE_ACTION_HINTS["polish"]
    if act == "free" and goal:
        hint = goal
    elif act == "free" and not goal:
        hint = OFFICE_ACTION_HINTS["polish"]

    parts = [
        "用户正在文字处理软件（WPS/Word）中编辑，下列【选中文本】来自当前选区。",
        hint,
    ]
    if goal and act != "free":
        parts.append(f"用户补充说明：{goal}")
    ref_list = reference_docs or []
    if ref_list:
        ref_chunks = []
        for label, body in ref_list:
            ref_chunks.append(f"【参考文档：{label}】\n{body}")
        parts.append(
            "下列【参考文档】由用户主动投喂，与本次写作相关。辅助时请结合其中设定、事实与风格；"
            "不要编造与参考明显矛盾的内容；若与当前选区无关可少引用。\n\n"
            + "\n\n---\n\n".join(ref_chunks)
        )
    if ctx:
        parts.append(
            "【附近上下文（仅供参考，勿大段复述）】\n"
            + ctx[:OFFICE_CONTEXT_MAX]
        )
    parts.append("【选中文本】\n" + passage[:OFFICE_PASSAGE_MAX])
    return "\n\n".join(parts)


def build_creative_prompt(doc: LoadedDocument, user_goal: Optional[str] = None) -> str:
    goal = (user_goal or "").strip()
    header = "你将作为创作助手阅读用户文档，并给出可直接采纳的建议。"
    if goal:
        header += f"\n用户目标：{goal}"
    return (
        header
        + "\n\n【文档内容（已截断）】\n"
        + doc.text.strip()
        + "\n\n请输出：\n"
        + "1) 作品一句话提炼\n"
        + "2) 3 条最具体可执行的改进建议\n"
        + "3) 给 2 段可直接复制粘贴的“替换/续写”文本（各 2-4 句）\n"
    )

